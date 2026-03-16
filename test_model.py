from mlx_lm import load, generate
from docx import Document
import json
import re
import os
from datetime import datetime

# --- 1. HÀM ĐỌC WORD ---
def read_legal_docx_universal(file_path):
    if not os.path.exists(file_path):
        return None
    doc = Document(file_path)
    lines = []
    
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip(): lines.append(f"[HEADER] {p.text.strip()}")
        for table in section.header.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                lines.append(f"[HEADER_TABLE] {' | '.join(dict.fromkeys(cells))}")

    for p in doc.paragraphs:
        if p.text.strip(): lines.append(p.text.strip())
    for t in doc.tables:
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells if c.text.strip()]
            lines.append(f"[TABLE] {' | '.join(dict.fromkeys(cells))}")

    for section in doc.sections:
        for p in section.footer.paragraphs:
            if p.text.strip(): lines.append(f"[FOOTER] {p.text.strip()}")
            
    raw_text = "\n".join(lines)
    
    raw_text = raw_text.replace('\xa0', ' ').replace('\t', ' ')
    raw_text = re.sub(r'["“”\']', '', raw_text)
    raw_text = re.sub(r'(?i)\s+(Căn cứ\s+)', r'\n\1', raw_text)
    raw_text = re.sub(r'(?i)\s+(Chương\s+[IVXLCDM]+)', r'\n\1', raw_text)
    raw_text = re.sub(r'(?i)(Điều\s+\d+[\.:]?)', r'\n\1', raw_text)
    raw_text = re.sub(r'\n+', '\n', raw_text)
    
    return raw_text[:15000]

# --- 2. HÀM FIX LỖI JSON ---
def auto_fix_json(json_str):
    json_str = re.sub(r'}\s*{', '}, {', json_str)
    json_str = re.sub(r',\s*([\]}])', r'\1', json_str)
    return json_str

# --- 3. CẤU HÌNH MODEL ---
MODEL_PATH = "mlx-community/Qwen2.5-3B-Instruct-4bit"
print("--- Đang tải Model MLX (GPU M1)... ---")
model, tokenizer = load(MODEL_PATH)

def run_universal_extraction(docx_name, json_name):
    raw_text = read_legal_docx_universal(docx_name)
    if not raw_text:
        print("❌ Lỗi: Không tìm thấy file đầu vào.")
        return

    # --- 4. PROMPT: DẠY AI PHÂN BIỆT RÕ 3 TRƯỜNG THUỘC TÍNH ---
    prompt = f"""<|im_start|>system
Bạn là AI chuyên trích xuất dữ liệu pháp luật. 
LỆNH BẮT BUỘC ĐỂ KHÔNG BỊ LẪN LỘN DỮ LIỆU:
1. PHÂN BIỆT RÕ 3 TRƯỜNG SAU (TUYỆT ĐỐI KHÔNG TẠO MẢNG [], CHỈ DÙNG CHUỖI):
   - 'co_quan_ban_hanh': Chỉ chứa tên Cơ Quan (VD: Bộ Tư pháp, Bộ Ngoại giao, Tòa án nhân dân tối cao). Không chứa tên người.
   - 'chuc_danh': Chỉ chứa Chức Vụ (VD: Phó Chánh án, Thứ trưởng). KHÔNG LẤY TÊN NGƯỜI, KHÔNG LẤY CƠ QUAN. Bỏ các từ thừa như "KT. Bộ trưởng".
   - 'nguoi_ky': Chỉ chứa Họ và Tên (VD: Phạm Quốc Hưng, Lê Thị Thu Hằng, Nguyễn Thanh Tịnh). KHÔNG LẤY CHỨC VỤ.
2. 'can_cu': Phải liệt kê đủ các dòng "Căn cứ..." và cả các dùng sau căn cứ là dòng "Bộ...".
3. 'dieu': Tách riêng nội dung Điều 1, Điều 10, Điều 17, Điều 2, Điều 3 thành các chuỗi dài. Không dùng mảng.

MẪU JSON BẮT BUỘC (Thay thế nội dung vào chỗ trống):
{{
  "thuoc_tinh": {{
    "so_ky_hieu": "<chuỗi>",
    "ngay_ban_hanh": "<chuỗi>",
    "loai_van_ban": "<chuỗi>",
    "co_quan_ban_hanh": "<chuỗi tên cơ quan, phân cách bằng dấu phẩy>",
    "dia_diem_ban_hanh": "<chuỗi>",
    "chuc_danh": "<chuỗi chức danh, phân cách bằng dấu phẩy>",
    "nguoi_ky": "<chuỗi tên người, phân cách bằng dấu phẩy>"
  }},
  "can_cu": [
    "<chuỗi 1>",
    "<chuỗi 2>"
  ],
  "chuong": [],
  "dieu": [
    {{
      "ten_dieu": "Điều 1",
      "noi_dung": "<chuỗi>"
    }},
    {{
      "ten_dieu": "Điều 10",
      "noi_dung": "<chuỗi>"
    }},
    {{
      "ten_dieu": "Điều 17",
      "noi_dung": "<chuỗi>"
    }},
    {{
      "ten_dieu": "Điều 2",
      "noi_dung": "<chuỗi>"
    }},
    {{
      "ten_dieu": "Điều 3",
      "noi_dung": "<chuỗi>"
    }}
  ]
}}
<|im_end|>
<|im_start|>user
VĂN BẢN:
{raw_text}

Trích xuất JSON ngay lập tức.<|im_end|>
<|im_start|>assistant
"""

    print("--- AI đang trích xuất dữ liệu... ---")
    response = generate(model, tokenizer, prompt=prompt, max_tokens=4000)

    try:
        start_idx = response.find('{')
        end_idx = response.rfind('}')
        if start_idx == -1 or end_idx == -1:
            raise ValueError("Không tìm thấy cấu trúc JSON.")
            
        clean_json = response[start_idx:end_idx+1]
        clean_json = auto_fix_json(clean_json)
        
        data = json.loads(clean_json, strict=False)
        
        # --- BỘ LỌC ÉP KIỂU BẰNG PYTHON (ÉP MẢNG VỀ CHUỖI) ---
        if "thuoc_tinh" in data:
            thuoc_tinh = data["thuoc_tinh"]
            # Nếu AI vẫn ngoan cố tạo mảng, Python sẽ gộp mảng thành chuỗi phân cách bằng dấu phẩy
            for key in ["co_quan_ban_hanh", "dia_diem_ban_hanh", "chuc_danh", "nguoi_ky"]:
                if key in thuoc_tinh and isinstance(thuoc_tinh[key], list):
                    # Lọc bỏ các giá trị rác bị lặp lại nhiều lần (như "KT. Bộ trưởng")
                    unique_vals = list(dict.fromkeys([str(item).strip() for item in thuoc_tinh[key]]))
                    thuoc_tinh[key] = ", ".join(unique_vals)
            data["thuoc_tinh"] = thuoc_tinh
        # ----------------------------------------------------

        data["metadata_he_thong"] = {
            "ngay_crawl": datetime.now().isoformat(),
            "source_url": ""
        }
        
        with open(json_name, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"✅ THÀNH CÔNG: Đã lưu file '{json_name}'")
        
    except json.decoder.JSONDecodeError as e:
        print(f"❌ Lỗi cú pháp JSON: {e}")
        err_idx = e.pos
        print(f"⚠️ Dữ liệu lỗi quanh khu vực này: \n...{clean_json[max(0, err_idx-50) : min(len(clean_json), err_idx+50)]}...")
        with open("raw_debug.txt", "w", encoding="utf-8") as f:
            f.write(clean_json)
    except Exception as e:
        print(f"❌ Lỗi hệ thống: {e}")

if __name__ == "__main__":
    run_universal_extraction("Tên file để cùng thư mục.docx", "output.json")
